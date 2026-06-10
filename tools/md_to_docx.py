"""Convert the project's Markdown docs to .docx via python-docx.

Handles the subset used in docs/: # / ## / ### headings, paragraphs,
- bullet lists, numbered lists, and | pipe | tables. The first H1 becomes the
document title. Keep authoring in Markdown (diffable, reviewable); regenerate
the .docx after edits.

Usage:  python tools/md_to_docx.py docs/BRINGUP_NOTEBOOK.md [out.docx]
"""

import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ACCENT = RGBColor(0x1D, 0x7A, 0x48)   # phosphor-green-adjacent, print-friendly
GREY = RGBColor(0x55, 0x5F, 0x66)


def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def style_doc(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for level, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
        st = doc.styles[level]
        st.font.name = "Cambria"
        st.font.size = Pt(size)
        st.font.color.rgb = ACCENT
        st.font.bold = True


def add_table(doc, rows):
    header, body = rows[0], [r for r in rows[1:] if not re.match(r"^[\s\-:|]+$", "|".join(r))]
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Light Grid Accent 3"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = cell_text
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9.5)
    for ri, row in enumerate(body, start=1):
        for ci in range(len(header)):
            cell = table.rows[ri].cells[ci]
            cell.text = row[ci] if ci < len(row) else ""
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
    doc.add_paragraph()


def convert(md_path, out_path):
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    style_doc(doc)

    i = 0
    first_h1_done = False
    while i < len(lines):
        line = lines[i]

        if is_table_row(line):
            rows = []
            while i < len(lines) and is_table_row(lines[i]):
                rows.append(split_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("# ") and not first_h1_done:
            first_h1_done = True
            p = doc.add_paragraph(stripped[2:])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.name = "Cambria"
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = ACCENT
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            # Merge hard-wrapped continuation lines into one paragraph.
            para = [stripped]
            while (i + 1 < len(lines) and lines[i + 1].strip()
                   and not lines[i + 1].strip().startswith(("#", "-", "|"))
                   and not re.match(r"^\d+\.\s", lines[i + 1].strip())):
                i += 1
                para.append(lines[i].strip())
            text = " ".join(para)
            p = doc.add_paragraph(text)
            # Document metadata block right under the title gets muted styling.
            if not any(s.startswith("#") for s in lines[:max(0, i - len(para))]) is False:
                pass
            if text.startswith(("Project:", "Repository:", "Period:", "Status:", "End of notebook")):
                for r in p.runs:
                    r.font.color.rgb = GREY
                    r.font.size = Pt(9.5)
        i += 1

    doc.save(out_path)
    print("wrote", out_path)


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else "docs/BRINGUP_NOTEBOOK.md"
    dst = sys.argv[2] if len(sys.argv) > 2 else re.sub(r"\.md$", ".docx", src)
    convert(src, dst)
